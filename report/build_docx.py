# -*- coding: utf-8 -*-
"""Build the CSE471 assignment report DOCX for Matir Poshara."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
SHOTS = ROOT / "screenshots"
DIA = ROOT / "diagrams"

TERRA = RGBColor(0x9C, 0x4A, 0x2A)
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def set_complex_font(style, name="Nirmala UI"):
    """Ensure Bengali (complex-script) glyphs render with a Bangla-capable font."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), name)


set_complex_font(normal)

for i, sz in [(1, 18), (2, 14), (3, 12)]:
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = TERRA
    st.font.bold = True
    set_complex_font(st)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def body(text, italic=False, size=11, after=6, color=None, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def figure(img, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(width))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(12)


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


# ======================================================================
# COVER PAGE
# ======================================================================
for _ in range(2):
    doc.add_paragraph()

logo = doc.add_paragraph()
logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = logo.add_run("🏺")
r.font.size = Pt(54)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Matir Poshara (মাটির পশরা)")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = TERRA

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("A Full-Stack E-Commerce Platform for Traditional Bangladeshi Pottery")
r.font.size = Pt(13)
r.font.color.rgb = DARK
r.italic = True

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Course Project Report")
r.font.size = Pt(14)
r.font.bold = True

# info table
info = doc.add_table(rows=0, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = "Table Grid"
rows = [
    ("Course", "CSE471 — [Course Title]"),
    ("Project Title", "Matir Poshara — Online Pottery Marketplace"),
    ("Group ID", "[Group ID]"),
    ("Group Members", "Ferdous Hasan Rahid — 2023100000546\n[Add additional group members here]"),
    ("GitHub Repository", "https://github.com/amnahid/matir-poshara-next"),
    ("Live Website", "[Deployment URL — e.g. https://matir-poshara.vercel.app]"),
    ("Submission Date", "9 June 2026"),
]
for k, v in rows:
    cells = info.add_row().cells
    cells[0].width = Inches(1.8)
    cells[1].width = Inches(4.2)
    set_cell_bg(cells[0], "F3E3D8")
    pk = cells[0].paragraphs[0]
    rk = pk.add_run(k)
    rk.bold = True
    rk.font.size = Pt(10.5)
    rk.font.color.rgb = TERRA
    for li, line in enumerate(v.split("\n")):
        pv = cells[1].paragraphs[0] if li == 0 else cells[1].add_paragraph()
        rv = pv.add_run(line)
        rv.font.size = Pt(10.5)

doc.add_page_break()

# ======================================================================
# 1. OBJECTIVES & SCOPE
# ======================================================================
h1("1. Objectives & Scope")
body(
    "Matir Poshara (“The Clay Marketplace”) is a full-stack e-commerce web application "
    "that brings Bangladesh’s traditional terracotta and clay-pottery artisans online. The platform "
    "lets customers browse handcrafted products by category, search the catalogue, view detailed product "
    "pages, place orders through a guest-friendly checkout, and manage their own accounts. A separate, "
    "password-protected admin panel gives the store owner full control over inventory, orders, artisans, "
    "and site-wide settings."
)
h2("1.1 Objectives")
for o in [
    "Build a responsive, bilingual (Bengali-first) storefront that showcases pottery products with images, pricing, discounts, and live stock status.",
    "Provide secure customer account management — registration, login, profile editing, and password change — using hashed credentials and signed session cookies.",
    "Deliver a complete admin CRUD system for products, including bulk CSV import/export and real-time inventory tracking.",
    "Implement an order-management workflow from cart to checkout to admin fulfilment with status tracking.",
    "Offer extensibility hooks: an admin-configurable live-chat widget and a RAG/agentic integration that keeps an external AI service in sync with the product database in real time.",
]:
    bullet(o)

h2("1.2 Scope")
body("The delivered system covers the following functional areas:")
for s in [
    "Customer storefront: home/landing, category browsing with filters, full-text product search, product detail, special-offers page, shopping cart, and checkout.",
    "Authentication: customer register/login/logout/account; localStorage-gated admin login.",
    "Admin panel: dashboard with KPIs, product CRUD + CSV import/export, order management, artisan management, and a settings page (live chat + RAG).",
    "Persistence: all entities stored in MongoDB Atlas; inventory and orders updated in real time.",
]:
    bullet(s)
body(
    "Out of scope for this iteration: online payment-gateway integration (orders are recorded as "
    "cash-on-delivery), multi-vendor seller accounts, and shipping-logistics tracking.",
    italic=True, color=GREY,
)

# ======================================================================
# 2. ARCHITECTURE
# ======================================================================
h1("2. System Architecture")
body(
    "The application follows a classic three-tier architecture — Client, Server, and Database — built on "
    "the Next.js 16 App Router. The same Next.js process serves both the React user interface (via Server "
    "and Client Components) and the backend REST API (via Route Handlers), while Mongoose connects the "
    "server tier to a MongoDB Atlas cluster."
)
figure(DIA / "architecture.png", "Figure 1: Three-tier architecture — Client → Server → Database, with external Live-Chat and RAG integrations.")
h2("2.1 Technology Stack")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for c, txt in zip(hdr, ["Layer", "Technologies"]):
    set_cell_bg(c, "9C4A2A")
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
stack = [
    ("Client", "React 19, Next.js 16 App Router, Tailwind CSS v4, Framer Motion, lucide-react icons"),
    ("Server", "Next.js Route Handlers (Node runtime), Server Components, Node crypto (scrypt + HMAC)"),
    ("Database", "MongoDB Atlas (replica set), Mongoose 9 ODM, database name “matir-poshra”"),
    ("Integrations", "Configurable live-chat iframe; RAG webhook push + MongoDB Change Streams"),
]
for layer, tech in stack:
    cells = tbl.add_row().cells
    set_cell_bg(cells[0], "F3E3D8")
    rr = cells[0].paragraphs[0].add_run(layer)
    rr.bold = True
    cells[1].paragraphs[0].add_run(tech)

h2("2.2 Request Flow")
body(
    "A typical request flows as follows: the browser renders a Server Component (e.g. the category page) "
    "whose data is fetched on the server directly through Mongoose, or it issues a client-side fetch to a "
    "Route Handler under /api. The handler validates the request, executes the Mongoose query against "
    "MongoDB Atlas over TLS, and returns JSON. Authenticated routes read the signed mp_session cookie; "
    "admin product mutations additionally fire a best-effort webhook to the configured RAG service so the "
    "external AI index stays current."
)

# ======================================================================
# 3. DATABASE SCHEMA
# ======================================================================
h1("3. Database Schema")
body(
    "Matir Poshara uses a document-oriented schema in MongoDB. Six collections model the domain: "
    "users, products, categories, artisans, orders, and a singleton settings document. Each document "
    "carries an auto-generated _id (ObjectId) primary key; selected fields enforce uniqueness."
)
figure(DIA / "er.png", "Figure 2: Entity–Relationship diagram of the six MongoDB collections, with primary, foreign, and unique keys.")

h2("3.1 Collection Definitions & Keys")
schema_tables = [
    ("users", [
        ("_id", "ObjectId", "Primary key (auto)"),
        ("name", "String", "Required"),
        ("email", "String", "Required, UNIQUE, lowercased"),
        ("phone / address", "String", "Optional profile fields"),
        ("passwordHash", "String", "scrypt salt:hash, required"),
        ("createdAt / updatedAt", "Date", "Timestamps"),
    ]),
    ("products", [
        ("_id", "ObjectId", "Primary key (auto)"),
        ("name / description", "String", "name required"),
        ("price / originalPrice", "Number", "price required; original drives discount %"),
        ("category", "String", "Reference → categories.name"),
        ("stock", "Number", "Inventory count (min 0)"),
        ("badge / image / icon", "String", "Display metadata"),
        ("isBestSelling", "Boolean", "Featured flag"),
    ]),
    ("orders", [
        ("_id", "ObjectId", "Primary key (auto)"),
        ("orderNumber", "String", "UNIQUE, format MP-YYYYMMDD-####"),
        ("customer", "Embedded", "{ name, phone, address }"),
        ("items[]", "Embedded array", "{ productId→products._id, name, price, qty }"),
        ("totalPrice", "Number", "Required"),
        ("status", "String (enum)", "pending|processing|shipped|delivered|cancelled"),
    ]),
    ("categories / artisans / settings", [
        ("categories", "Document", "name, productCount, image, icon"),
        ("artisans", "Document", "name, village, experience, story, image"),
        ("settings", "Singleton", "key=“global” (UNIQUE); live-chat + RAG config"),
    ]),
]
for title, rows in schema_tables:
    body(title, bold=True, after=2, color=TERRA)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for c, txt in zip(t.rows[0].cells, ["Field", "Type", "Constraint / Notes"]):
        set_cell_bg(c, "9C4A2A")
        rr = c.paragraphs[0].add_run(txt)
        rr.bold = True
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rr.font.size = Pt(10)
    for f, ty, note in rows:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(f).font.size = Pt(10)
        cells[1].paragraphs[0].add_run(ty).font.size = Pt(10)
        cells[2].paragraphs[0].add_run(note).font.size = Pt(10)
    doc.add_paragraph()

body(
    "Relationships: a Category has many Products (Product.category references Category.name); a Product "
    "appears in many Order items (items[].productId references Product._id); a User may place many Orders. "
    "Because MongoDB is non-relational, these are application-level references rather than enforced foreign "
    "keys, which keeps reads fast and documents self-contained."
)

# ======================================================================
# 4. FEATURES
# ======================================================================
h1("4. Features")
body(
    "This section walks through the implemented features with screenshots captured from the running "
    "application and the live MongoDB Atlas database (28 seeded products, a demo customer, and demo orders)."
)

features = [
    ("4.1 Storefront & Landing Page", "01_home",
     "Figure 3: Landing page.",
     "The home page presents the brand hero, featured categories, best-selling products, and artisan "
     "stories. The sticky header carries a search bar, wishlist/cart icons, and account links; the navbar "
     "lists all six product categories plus an Offers shortcut."),
    ("4.2 Customer Registration", "02_register",
     "Figure 4: Customer registration form.",
     "New customers register with name, email, phone, address, and password. The server checks for "
     "duplicate emails, hashes the password with scrypt, persists the user to MongoDB, and immediately "
     "issues a signed session cookie."),
    ("4.3 Customer Login", "03_login",
     "Figure 5: Login page.",
     "Returning customers log in with email and password. Credentials are verified using a constant-time "
     "comparison against the stored scrypt hash, and a 30-day HMAC-signed httpOnly cookie is set."),
    ("4.4 Account Management", "04_account",
     "Figure 6: Logged-in customer account page.",
     "Authenticated users can view and update their profile (name, phone, address) and change their "
     "password. The page is protected — unauthenticated visitors are redirected to the login page."),
    ("4.5 Search", "05_search",
     "Figure 7: Search results for the query “মাটি”.",
     "A dedicated /api/search endpoint performs a case-insensitive match across product names and "
     "descriptions and returns matching products, which are rendered with live stock status."),
    ("4.6 Category Browsing & Filters", "06_category",
     "Figure 8: Category page with sidebar filters.",
     "Each category page lists its products and offers price-range and rating filters via query "
     "parameters, resolved on the server through Mongoose queries."),
    ("4.7 Product Detail", "07_product",
     "Figure 9: Product detail page.",
     "The product page shows the image/icon, price with any discount, rating, description, and a "
     "stock-aware add-to-cart button that disables itself and shows “out of stock” when inventory is zero."),
    ("4.8 Special Offers", "08_offers",
     "Figure 10: Offers page listing discounted products.",
     "The offers page queries products whose originalPrice exceeds the current price using a MongoDB "
     "$expr comparison, surfacing every active discount in one place."),
    ("4.9 Checkout", "09_checkout",
     "Figure 11: Checkout page with cart summary.",
     "Cart state lives in localStorage; at checkout the customer supplies delivery details and the order "
     "is posted to /api/orders, which generates a unique order number and stores the order in MongoDB."),
    ("4.10 Admin Login", "10_admin_login",
     "Figure 12: Admin panel login.",
     "The admin area is gated by a separate login. On success an admin_auth flag is stored and the "
     "dashboard becomes accessible; every admin route checks this flag before rendering."),
    ("4.11 Admin Dashboard", "11_admin_dashboard",
     "Figure 13: Admin dashboard with live KPIs.",
     "The dashboard aggregates real figures from the database — total products, orders, customers, and "
     "revenue — alongside the most recent orders and quick-action shortcuts."),
    ("4.12 Inventory Management (Read / Update / Delete)", "12_admin_products",
     "Figure 14: Product inventory table with stock badges and CSV import/export.",
     "Admins see every product with its exact stock count and an in-stock / out-of-stock badge "
     "(customers see only the badge). Each row can be edited or deleted, and the whole catalogue can be "
     "imported or exported as CSV."),
    ("4.13 Add Product (Create)", "13_admin_product_form",
     "Figure 15: Add / edit product modal.",
     "The create/edit form captures name, description, category, badge, price, original price, stock, "
     "image URL, icon, and a best-selling flag. New products without an image fall back to a placeholder, "
     "preserving the layout."),
    ("4.14 Order Management", "14_admin_orders",
     "Figure 16: Order management with status control.",
     "All customer orders are listed with order number, customer, items, total, and status. Admins can "
     "advance each order through its lifecycle (pending → processing → shipped → delivered)."),
    ("4.15 Settings — Live Chat & RAG Integration", "16_admin_settings",
     "Figure 17: Settings page for live-chat and RAG integration.",
     "Admins can enable a floating live-chat widget by pasting any embed URL, and connect an external "
     "RAG/agentic service via a webhook URL and write-only API key. A “sync now” action pushes the entire "
     "catalogue, and the page documents the direct MongoDB Change-Streams option."),
]
for title, img, cap, desc in features:
    h2(title)
    body(desc)
    figure(SHOTS / f"{img}.png", cap, width=5.1)

# ======================================================================
# 5. VALIDATION & SECURITY
# ======================================================================
h1("5. Validation & Security")
body("Secure-coding practices were applied across the stack:")
sec = [
    ("Password hashing", "User passwords are never stored in plaintext. They are hashed with Node’s built-in scrypt (a memory-hard KDF) using a per-user random salt, stored as salt:hash."),
    ("Constant-time verification", "Login compares hashes with crypto.timingSafeEqual to prevent timing attacks."),
    ("Signed sessions", "Sessions use an HMAC-signed, httpOnly cookie (mp_session) so the token cannot be read by JavaScript or forged client-side."),
    ("Secret management", "The MongoDB connection string and AUTH_SECRET live in .env.local, which is git-ignored and never shipped to the client."),
    ("Write-only secrets in API", "The RAG API key is write-only — admin GET responses return only a boolean ragApiKeySet, never the key value. The public /api/settings endpoint exposes only live-chat fields and never any RAG secret."),
    ("Input validation", "API routes validate required fields (e.g. registration rejects missing/duplicate emails; order creation rejects empty carts) and return appropriate 4xx status codes."),
    ("Server-side stock truth", "Inventory counts are computed and enforced on the server; customers receive only an in/out-of-stock boolean, never the raw count."),
    ("Schema constraints", "Mongoose schemas enforce required fields, unique indexes (email, orderNumber, settings.key), enums (order status), and minimums (stock ≥ 0)."),
]
for k, v in sec:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(k + ": ")
    r.bold = True
    p.add_run(v)
    p.paragraph_format.space_after = Pt(3)
body(
    "Known limitation & future work: the /api/admin/* routes are currently gated only by the client-side "
    "admin_auth flag and are not yet protected by server-side session checks. Hardening these endpoints with "
    "a real admin session (mirroring the customer cookie scheme) is the top recommended next step before "
    "production deployment.",
    italic=True, color=GREY,
)

# ======================================================================
# 6. TESTING
# ======================================================================
h1("6. Testing")
body(
    "The application was tested with a mix of valid, invalid, and edge-case scenarios covering "
    "authentication, CRUD, search, and ordering. The table below summarises the core test cases."
)
test_rows = [
    ("TC-01", "Customer registration (valid)", "Submit unique email + valid fields", "200; user persisted in DB, session cookie set", "Pass"),
    ("TC-02", "Duplicate registration (invalid)", "Register with an already-used email", "Rejected with “email already exists”, no duplicate created", "Pass"),
    ("TC-03", "Login with wrong password (invalid)", "Correct email, incorrect password", "Login refused; no session issued", "Pass"),
    ("TC-04", "Product search", "Search query “মাটি”", "Only matching products returned with stock status", "Pass"),
    ("TC-05", "Out-of-stock product (edge case)", "Open a product whose stock = 0", "Add-to-cart disabled; “out of stock” shown to customer", "Pass"),
    ("TC-06", "Admin create + CSV import", "Add a product / import CSV rows", "New products appear in inventory table and DB", "Pass"),
    ("TC-07", "Place order (valid)", "Checkout with items in cart", "201; unique MP-YYYYMMDD order created, visible in admin", "Pass"),
    ("TC-08", "Empty-cart checkout (edge case)", "POST order with no items", "400 Invalid order data; no order created", "Pass"),
]
tt = doc.add_table(rows=1, cols=5)
tt.style = "Table Grid"
for c, txt in zip(tt.rows[0].cells, ["ID", "Test Case", "Input / Action", "Expected Result", "Status"]):
    set_cell_bg(c, "9C4A2A")
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(9.5)
for tid, name, inp, exp, status in test_rows:
    cells = tt.add_row().cells
    vals = [tid, name, inp, exp, status]
    for c, v in zip(cells, vals):
        rr = c.paragraphs[0].add_run(v)
        rr.font.size = Pt(9.5)
        if v == "Pass":
            rr.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            rr.bold = True

# ======================================================================
# 7. TASK DIVISION
# ======================================================================
h1("7. Task Division")
body(
    "The contribution of each group member across the project areas is summarised below. (Replace the "
    "placeholder rows with your actual teammates and adjust the split as appropriate.)"
)
task_tbl = doc.add_table(rows=1, cols=6)
task_tbl.style = "Table Grid"
for c, txt in zip(task_tbl.rows[0].cells, ["Member (ID)", "Frontend", "Backend", "Database", "Documentation", "Testing"]):
    set_cell_bg(c, "9C4A2A")
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(9.5)
task_data = [
    ("Ferdous Hasan Rahid (2023100000546)", "✓", "✓", "✓", "✓", "✓"),
    ("[Member 2 — ID]", "", "", "", "", ""),
    ("[Member 3 — ID]", "", "", "", "", ""),
    ("[Member 4 — ID]", "", "", "", "", ""),
]
for row in task_data:
    cells = task_tbl.add_row().cells
    for c, v in zip(cells, row):
        rr = c.paragraphs[0].add_run(v)
        rr.font.size = Pt(10)
        if v == "✓":
            rr.font.color.rgb = TERRA
            rr.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ======================================================================
# 8. REFLECTION
# ======================================================================
h1("8. Reflection")
h2("8.1 Lessons Learned")
for l in [
    "Working with the Next.js 16 App Router required adapting to its newer conventions — params and searchParams are asynchronous (Promises) and must be awaited — which reinforced the value of reading current framework documentation rather than relying on older patterns.",
    "A document database (MongoDB) made it natural to embed order items and a singleton settings document, simplifying reads at the cost of application-managed references.",
    "Security is easier when built in from the start: using Node’s built-in crypto for scrypt hashing and HMAC sessions avoided extra dependencies while keeping credentials safe.",
    "Designing the inventory model so that the server is the single source of truth — exposing only an in/out-of-stock boolean to customers — cleanly separated admin and customer concerns.",
]:
    bullet(l)
h2("8.2 Challenges Faced")
for c in [
    "Rendering Bengali (Bangla) text consistently across the UI and ensuring correct fonts and encoding throughout the stack.",
    "Correlating inventory between the admin view (exact counts, editable) and the customer view (status only) without leaking data.",
    "Building dependency-free CSV import/export with correct quoting and a UTF-8 BOM so Bengali product names open correctly in spreadsheet software.",
    "Keeping an external RAG service in sync in real time, which led to a dual approach: best-effort webhook pushes plus the option of MongoDB Change Streams.",
]:
    bullet(c)

# ======================================================================
# 9. REFERENCES
# ======================================================================
h1("9. References")
refs = [
    "MongoDB, Inc. (2026) MongoDB Atlas Documentation. Available at: https://www.mongodb.com/docs/atlas/ (Accessed: 9 June 2026).",
    "Mongoose (2026) Mongoose ODM Documentation. Available at: https://mongoosejs.com/docs/ (Accessed: 9 June 2026).",
    "Vercel (2026) Next.js Documentation. Available at: https://nextjs.org/docs (Accessed: 9 June 2026).",
    "Meta Open Source (2026) React Documentation. Available at: https://react.dev/ (Accessed: 9 June 2026).",
    "Tailwind Labs (2026) Tailwind CSS Documentation. Available at: https://tailwindcss.com/docs (Accessed: 9 June 2026).",
    "Node.js (2026) Crypto | Node.js Documentation. Available at: https://nodejs.org/api/crypto.html (Accessed: 9 June 2026).",
    "OWASP Foundation (2021) OWASP Top Ten. Available at: https://owasp.org/www-project-top-ten/ (Accessed: 9 June 2026).",
]
for r in refs:
    p = doc.add_paragraph()
    p.add_run(r).font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(6)

out = ROOT / "GroupID_CSE471_Assignment.docx"
doc.save(str(out))
print("SAVED", out)
